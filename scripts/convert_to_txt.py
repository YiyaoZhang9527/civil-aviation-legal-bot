#!/usr/bin/env python3
"""将 data/民航法律法规原文 中的 PDF/DOCX/DOC/TXT 统一转为纯文本到 data/法律数据。

用法:
    python scripts/convert_to_txt.py
    python scripts/convert_to_txt.py --dry-run   # 只打印计划，不写文件
"""

from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

# ── 配置 ──────────────────────────────────────────────
PROJECT = Path(__file__).resolve().parent.parent
SRC_DIR = PROJECT / "data" / "民航法律法规原文"
DST_DIR = PROJECT / "data" / "法律数据"

# 支持的文件后缀 → 处理方式
HANDLERS = {
    ".txt": "txt",
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "doc",
}


# ── 辅助函数 ──────────────────────────────────────────


def _table_to_markdown(data: list[list[str | None]]) -> str:
    """将二维数组转为 Markdown 表格。空单元格用空字符串。"""
    if not data:
        return ""
    ncols = max(len(row) for row in data)
    lines: list[str] = []
    for i, row in enumerate(data):
        cells = [str(c).replace("\n", " ").strip() if c else "" for c in row]
        # 补齐列数
        cells += [""] * (ncols - len(cells))
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("|" + "|".join("---" for _ in range(ncols)) + "|")
    return "\n".join(lines)


# ── 转换函数（一个函数只做一件事） ────────────────────


def convert_txt(src: Path) -> str:
    """TXT 直接读取，统一换行。"""
    raw = src.read_bytes()
    # 尝试 UTF-8，回退 GBK
    for enc in ("utf-8", "gbk", "gb18030"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")


def convert_pdf(src: Path) -> str:
    """用 PyMuPDF 提取 PDF 文本，表格保留为 Markdown 格式。扫描件自动 OCR。"""
    import fitz  # PyMuPDF

    doc = fitz.open(str(src))
    # 预检：是否为纯扫描件（无文本层）
    has_any_text = any(page.get_text().strip() for page in doc)
    if not has_any_text and doc.page_count > 0:
        doc.close()
        return _ocr_pdf(src)

    pages: list[str] = []
    for page in doc:
        # 1. 检测表格及其 bbox
        table_finder = page.find_tables()
        table_bboxes = [t.bbox for t in table_finder.tables]
        table_mds = [_table_to_markdown(t.extract()) for t in table_finder.tables]

        # 2. 构建表格插入点：按 y 坐标排序
        table_items = sorted(
            [(bboxes[1], md) for bboxes, md in zip(table_bboxes, table_mds) if md],
            key=lambda x: x[0],
        )

        if not table_items:
            text = page.get_text()
            if text.strip():
                pages.append(text)
            continue

        # 3. 获取文本块，跳过表格覆盖区域
        blocks = page.get_text("blocks")
        content_parts: list[tuple[float, str]] = []  # (y_pos, text)

        for block in blocks:
            # block: (x0, y0, x1, y1, text, block_no, block_type)
            if block[6] != 0:  # 非文本块
                continue
            bbox = (block[0], block[1], block[2], block[3])
            if _bbox_overlaps_tables(bbox, table_bboxes):
                continue
            text = block[4]
            if isinstance(text, bytes):
                text = text.decode("utf-8", errors="replace")
            if text.strip():
                content_parts.append((block[1], text.strip()))

        # 4. 插入表格到对应 y 位置
        for y, md in table_items:
            content_parts.append((y, "\n" + md + "\n"))

        # 5. 按 y 坐标排序后合并
        content_parts.sort(key=lambda x: x[0])
        page_text = "\n\n".join(t for _, t in content_parts)
        if page_text.strip():
            pages.append(page_text)

    doc.close()
    return "\n".join(pages)


def _ocr_pdf(src: Path) -> str:
    """对扫描件 PDF 用 Tesseract OCR 逐页识别。"""
    import fitz
    import pytesseract
    from PIL import Image
    import io

    doc = fitz.open(str(src))
    pages: list[str] = []
    for page in doc:
        pix = page.get_pixmap(dpi=300)
        img = Image.open(io.BytesIO(pix.tobytes('png')))
        text = pytesseract.image_to_string(img, lang='chi_sim')
        if text.strip():
            pages.append(text)
    doc.close()
    return "\n".join(pages)


def _bbox_overlaps_tables(bbox: tuple, table_bboxes: list) -> bool:
    """判断文本块 bbox 是否与任一表格 bbox 有重叠。"""
    x0, y0, x1, y1 = bbox
    for tb in table_bboxes:
        tx0, ty0, tx1, ty1 = tb
        # 矩形相交判定
        if x0 < tx1 and x1 > tx0 and y0 < ty1 and y1 > ty0:
            return True
    return False


def convert_docx(src: Path) -> str:
    """用 python-docx 提取 DOCX 文本，表格保留为 Markdown 格式。"""
    from docx import Document

    doc = Document(str(src))
    parts: list[str] = []

    # 按文档顺序遍历 body 子元素
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            # 提取段落文本
            texts = []
            for node in child.iter():
                if node.tag.endswith("}t") or node.tag == "t":
                    if node.text:
                        texts.append(node.text)
            line = "".join(texts).strip()
            if line:
                parts.append(line)
        elif tag == "tbl":
            # 提取表格为 Markdown
            rows = []
            for tr in child.iter():
                if tr.tag.endswith("}tr") or tr.tag == "tr":
                    cells = []
                    for tc in tr:
                        if tc.tag.endswith("}tc") or tc.tag == "tc":
                            cell_texts = []
                            for t_node in tc.iter():
                                if t_node.tag.endswith("}t") or t_node.tag == "t":
                                    if t_node.text:
                                        cell_texts.append(t_node.text)
                            cells.append(" ".join(cell_texts).strip())
                    if cells:
                        rows.append(cells)
            md = _table_to_markdown(rows)
            if md:
                parts.append(md)

    return "\n".join(parts)


def convert_doc(src: Path) -> str:
    """将 .doc 转为 txt。优先 python-docx，回退 LibreOffice/macOS textutil。"""
    # 方案 1: python-docx（能处理部分 .doc 文件）
    try:
        from docx import Document
        doc = Document(str(src))
        parts: list[str] = []
        for child in doc.element.body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                texts = []
                for node in child.iter():
                    if node.tag.endswith("}t") or node.tag == "t":
                        if node.text:
                            texts.append(node.text)
                line = "".join(texts).strip()
                if line:
                    parts.append(line)
            elif tag == "tbl":
                rows = []
                for tr in child.iter():
                    if tr.tag.endswith("}tr") or tr.tag == "tr":
                        cells = []
                        for tc in tr:
                            if tc.tag.endswith("}tc") or tc.tag == "tc":
                                cell_texts = []
                                for t_node in tc.iter():
                                    if t_node.tag.endswith("}t") or t_node.tag == "t":
                                        if t_node.text:
                                            cell_texts.append(t_node.text)
                                cells.append(" ".join(cell_texts).strip())
                        if cells:
                            rows.append(cells)
                md = _table_to_markdown(rows)
                if md:
                    parts.append(md)
        text = "\n".join(parts)
        if text.strip():
            return text
    except Exception:
        pass

    # 方案 2: LibreOffice（Linux / macOS 通用）
    for cmd in ("libreoffice", "soffice"):
        if _command_exists(cmd):
            return _convert_with_office(cmd, src)

    # 方案 3: macOS textutil（仅 macOS）
    import platform
    if platform.system() == "Darwin":
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp_path = tmp.name
        try:
            subprocess.run(
                ["textutil", "-convert", "txt", "-output", tmp_path, str(src)],
                check=True,
                capture_output=True,
            )
            return Path(tmp_path).read_text("utf-8")
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    raise RuntimeError(f"无法转换 .doc 文件：需要 python-docx、LibreOffice 或 macOS textutil: {src}")


def _command_exists(cmd: str) -> bool:
    """检查命令是否可用。"""
    import shutil
    return shutil.which(cmd) is not None


def _convert_with_office(cmd: str, src: Path) -> str:
    """用 LibreOffice 转换 .doc 为 txt。"""
    with tempfile.TemporaryDirectory() as tmp_dir:
        subprocess.run(
            [cmd, "--headless", "--convert-to", "txt:Text", "--outdir", tmp_dir, str(src)],
            check=True,
            capture_output=True,
            timeout=60,
        )
        out = Path(tmp_dir) / f"{src.stem}.txt"
        return out.read_text("utf-8") if out.exists() else ""


CONVERTERS = {
    "txt": convert_txt,
    "pdf": convert_pdf,
    "docx": convert_docx,
    "doc": convert_doc,
}


def clean_text(text: str) -> str:
    """最小化清洗：去掉连续空行（保留段落间隔）、全角空格规范化。"""
    lines = text.replace("\u3000", " ").splitlines()
    # 去掉每行首尾空白，保留空行作为段落分隔
    cleaned = [line.strip() for line in lines]
    # 合并连续空行为单个空行
    result: list[str] = []
    prev_empty = False
    for line in cleaned:
        if not line:
            if not prev_empty:
                result.append("")
            prev_empty = True
        else:
            result.append(line)
            prev_empty = False
    # 去掉首尾空行
    while result and not result[0]:
        result.pop(0)
    while result and not result[-1]:
        result.pop()
    return "\n".join(result) + "\n"


def output_stem(src_name: str) -> str:
    """生成干净的输出文件名（去掉书名号等特殊字符）。"""
    name = src_name
    for ch in "《》«»":
        name = name.replace(ch, "")
    return name


# ── 主流程 ────────────────────────────────────────────

def main(dry_run: bool = False) -> None:
    if not SRC_DIR.exists():
        print(f"源目录不存在: {SRC_DIR}", file=sys.stderr)
        sys.exit(1)

    DST_DIR.mkdir(parents=True, exist_ok=True)

    files = sorted(SRC_DIR.iterdir())
    stats = {"ok": 0, "skip": 0, "fail": 0}

    for src in files:
        if not src.is_file():
            continue

        ext = src.suffix.lower()
        kind = HANDLERS.get(ext)
        if kind is None:
            print(f"  跳过（不支持的格式）: {src.name}")
            stats["skip"] += 1
            continue

        dst_name = output_stem(src.stem) + ".txt"
        dst = DST_DIR / dst_name

        # 避免同名覆盖：如果已存在则加后缀
        if dst.exists():
            base = dst.stem
            idx = 1
            while dst.exists():
                dst = DST_DIR / f"{base}_{idx}.txt"
                idx += 1

        if dry_run:
            print(f"  {src.name} → {dst.name}")
            stats["ok"] += 1
            continue

        try:
            text = CONVERTERS[kind](src)
            text = clean_text(text)
            if not text.strip():
                print(f"  跳过（内容为空）: {src.name}")
                stats["skip"] += 1
                continue
            dst.write_text(text, encoding="utf-8")
            print(f"  ✓ {src.name} → {dst.name} ({len(text)} 字符)")
            stats["ok"] += 1
        except Exception as e:
            print(f"  ✗ {src.name}: {type(e).__name__}: {e}", file=sys.stderr)
            stats["fail"] += 1

    print(f"\n完成: 成功 {stats['ok']}, 跳过 {stats['skip']}, 失败 {stats['fail']}")


if __name__ == "__main__":
    dry_run = "--dry-run" in sys.argv
    main(dry_run=dry_run)
